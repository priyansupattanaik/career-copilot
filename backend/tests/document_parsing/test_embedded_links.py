from __future__ import annotations

import asyncio
import io
from types import SimpleNamespace

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas

from app.api.schemas import ProfileResumeExtractResult
from app.features.document_parsing.parsing.text_extract import (
    DOCX_MIME,
    PDF_MIME,
    extract_text,
)
from app.features.profile.agent import pipeline as profile_pipeline
from app.features.profile.agent.deterministic import build_profile_draft


def _docx_with_embedded_links() -> bytes:
    document = Document()
    document.add_heading("Priyansu Pattanaik", level=1)
    document.add_paragraph("Backend engineer building reliable recruitment systems with Python and FastAPI.")
    paragraph = document.add_paragraph("Profiles: ")
    for label, url in (
        ("LinkedIn", "https://www.linkedin.com/in/priyansu-pattanaik"),
        ("GitHub", "https://github.com/priyansu-pattanaik"),
        ("Portfolio", "https://priyansu.example.dev"),
    ):
        relationship_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = label
        run.append(text)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        paragraph.add_run("  ")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_with_embedded_link() -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 760, "Priyansu Pattanaik")
    pdf.drawString(72, 738, "Backend engineer building reliable recruitment systems with Python and FastAPI.")
    pdf.drawString(72, 716, "Portfolio")
    pdf.linkURL("https://priyansu.example.dev", (72, 712, 132, 728), relative=0)
    pdf.drawString(72, 694, "Experience: Built document parsing and semantic profile extraction pipelines.")
    pdf.drawString(72, 672, "Education: B.Tech in Electronics and Telecommunication Engineering.")
    pdf.save()
    return buffer.getvalue()


def test_docx_extraction_preserves_embedded_hyperlink_targets() -> None:
    extracted = extract_text(_docx_with_embedded_links(), DOCX_MIME)

    assert "LinkedIn https://www.linkedin.com/in/priyansu-pattanaik" in extracted
    assert "GitHub https://github.com/priyansu-pattanaik" in extracted
    assert "Portfolio https://priyansu.example.dev" in extracted


def test_pdf_extraction_preserves_link_annotation_targets() -> None:
    extracted = extract_text(_pdf_with_embedded_link(), PDF_MIME)

    # The annotation rect overlaps the visible "Portfolio" text, so the label
    # is recovered geometrically and the link is classified semantically.
    assert "Portfolio https://priyansu.example.dev" in extracted
    draft = build_profile_draft(extracted, {})
    assert draft["links"] == [
        {
            "link_type": "portfolio",
            "url": "https://priyansu.example.dev",
            "label": "Portfolio",
            "selected": True,
        }
    ]


def test_deterministic_links_use_labels_and_portfolio_hosts() -> None:
    text = (
        "Priyansu Pattanaik\n"
        "Backend engineer building reliable recruitment systems with Python and FastAPI.\n"
        "LinkedIn https://www.linkedin.com/in/priyansu-pattanaik\n"
        "GitHub https://github.com/priyansu-pattanaik\n"
        "My portfolio site https://priyansu.vercel.app\n"
        "Blog https://medium.com/@priyansu\n"
        "Experience: Built document parsing and semantic profile extraction pipelines.\n"
        "Education: B.Tech in Electronics and Telecommunication Engineering.\n"
        "Skills: Python, FastAPI, PostgreSQL, Docker.\n"
        "This paragraph exists so deterministic minimum-length checks pass on every run.\n"
    )
    draft = build_profile_draft(text, {})
    by_url = {link["url"]: link for link in draft["links"]}

    assert by_url["https://www.linkedin.com/in/priyansu-pattanaik"]["link_type"] == "linkedin"
    assert by_url["https://www.linkedin.com/in/priyansu-pattanaik"]["label"] == "LinkedIn"
    assert by_url["https://github.com/priyansu-pattanaik"]["link_type"] == "github"
    assert by_url["https://priyansu.vercel.app"]["link_type"] == "portfolio"
    assert by_url["https://priyansu.vercel.app"]["label"] == "My portfolio site"
    assert by_url["https://medium.com/@priyansu"]["link_type"] == "website"
    assert by_url["https://medium.com/@priyansu"]["label"] == "Blog"


def test_unlabeled_portfolio_host_is_classified_without_a_label() -> None:
    text = (
        "Priyansu Pattanaik\n"
        "Backend engineer building reliable recruitment systems with Python and FastAPI.\n"
        "https://priyansu.github.io\n"
        "Experience: Built document parsing and semantic profile extraction pipelines.\n"
        "Education: B.Tech in Electronics and Telecommunication Engineering.\n"
        "Skills: Python, FastAPI, PostgreSQL, Docker.\n"
        "This paragraph exists so deterministic minimum-length checks pass on every run.\n"
    )
    draft = build_profile_draft(text, {})
    assert draft["links"] == [
        {
            "link_type": "portfolio",
            "url": "https://priyansu.github.io",
            "label": None,
            "selected": True,
        }
    ]


def test_profile_agent_receives_embedded_targets_and_keeps_semantic_link_types(monkeypatch) -> None:
    extracted = extract_text(_docx_with_embedded_links(), DOCX_MIME)

    async def fake_generate(_settings, *, user_payload, **_kwargs):
        assert "https://www.linkedin.com/in/priyansu-pattanaik" in user_payload["resume_plain_text"]
        assert "https://github.com/priyansu-pattanaik" in user_payload["resume_plain_text"]
        assert "https://priyansu.example.dev" in user_payload["resume_plain_text"]
        return (
            ProfileResumeExtractResult.model_validate(
                {
                    "profile": {"full_name": "Priyansu Pattanaik", "current_role": "Backend engineer"},
                    "skills": ["Python", "FastAPI"],
                    "links": [
                        {
                            "link_type": "linkedin",
                            "url": "https://www.linkedin.com/in/priyansu-pattanaik",
                            "label": "LinkedIn",
                        },
                        {
                            "link_type": "github",
                            "url": "https://github.com/priyansu-pattanaik",
                            "label": "GitHub",
                        },
                        {
                            "link_type": "portfolio",
                            "url": "https://priyansu.example.dev",
                            "label": "Portfolio",
                        },
                    ],
                }
            ),
            "test-agent",
        )

    monkeypatch.setattr(profile_pipeline, "generate_structured_with_failover", fake_generate)

    draft = asyncio.run(
        profile_pipeline.build_profile_draft_enriched(extracted, {}, SimpleNamespace())
    )

    assert [(link["link_type"], link["url"]) for link in draft["links"]] == [
        ("linkedin", "https://www.linkedin.com/in/priyansu-pattanaik"),
        ("github", "https://github.com/priyansu-pattanaik"),
        ("portfolio", "https://priyansu.example.dev"),
    ]
    assert draft["meta"]["agent"] == "profile_fill"
    assert draft["meta"]["provider"] == "test-agent"
