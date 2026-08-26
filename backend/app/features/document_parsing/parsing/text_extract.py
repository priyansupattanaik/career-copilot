
from __future__ import annotations

import io
import logging
import re
from collections.abc import Callable, Iterable
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.core.errors import ApiError

logger = logging.getLogger(__name__)
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# Minimum usable extracted text length (chars). Applied on every accept path.
MIN_PDF_TEXT_CHARS = 200
MIN_DOCX_TEXT_CHARS = 80


def _normalize_extracted_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00ad", "").replace("\ufeff", "")
    lines: list[str] = []
    for raw in text.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if line.startswith("#"):
            line = re.sub(r"^#+\s*", "", line).strip()
        lines.append(line)
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if not line:
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")
            continue
        blank_run = 0
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def _docx_paragraph_text(document: Document) -> list[str]:
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            for cell in cells:
                for part in cell.splitlines():
                    part = part.strip()
                    if part:
                        lines.append(part)
    return lines


def _iter_docx_paragraphs(document: DocumentObject) -> Iterable[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _docx_embedded_links(document: DocumentObject) -> list[str]:
    """Return visible hyperlink labels together with their relationship targets."""
    links: list[str] = []
    seen: set[str] = set()
    for paragraph in _iter_docx_paragraphs(document):
        for hyperlink in paragraph._p.iter(qn("w:hyperlink")):
            relationship_id = hyperlink.get(qn("r:id"))
            if not relationship_id:
                continue
            relationship = paragraph.part.rels.get(relationship_id)
            target = str(getattr(relationship, "target_ref", "") or "").strip()
            if not target or target.casefold() in seen:
                continue
            label = "".join(node.text or "" for node in hyperlink.iter(qn("w:t"))).strip()
            links.append(f"{label} {target}".strip())
            seen.add(target.casefold())
    return links


def _make_text_run_collector(
    runs: list[tuple[str, float, float]],
) -> Callable[[str, Any, Any, Any, Any], None]:
    """Build a pypdf visitor_text callback appending positioned runs to ``runs``.

    The factory binds the list at call time so the closure never references a
    loop variable; each page gets an independent collector.
    """

    def _visit(text: str, cm: Any, tm: Any, font_dict: Any, font_size: Any) -> None:
        stripped = (text or "").strip()
        if stripped:
            runs.append((stripped, float(tm[4]), float(tm[5])))

    return _visit


def _pdf_embedded_links(content: bytes) -> list[str]:
    """Return external URI actions stored in PDF link annotations.

    When the annotation rectangle overlaps visible text, that text is kept as
    a label ("Portfolio https://…") so downstream parsing can classify the
    link semantically instead of guessing from the host alone. Label recovery
    uses pypdf's visitor text coordinates, so it works on the always-available
    backend without optional extras.
    """
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception:
        return []
    links: list[str] = []
    seen: set[str] = set()
    for page in reader.pages:
        rects: list[tuple[str, tuple[float, float, float, float] | None]] = []
        for annotation_ref in page.get("/Annots") or []:
            try:
                annotation = annotation_ref.get_object()
                if annotation.get("/Subtype") != "/Link":
                    continue
                action = annotation.get("/A")
                action = action.get_object() if hasattr(action, "get_object") else action
                target = str((action or {}).get("/URI") or "").strip()
                if not target or target.casefold() in seen:
                    continue
                rect: tuple[float, float, float, float] | None = None
                raw_rect = annotation.get("/Rect")
                if raw_rect:
                    try:
                        rx0, ry0, rx1, ry1 = (float(v) for v in raw_rect)
                        rect = (min(rx0, rx1), min(ry0, ry1), max(rx0, rx1), max(ry0, ry1))
                    except (TypeError, ValueError):
                        rect = None
                rects.append((target, rect))
                seen.add(target.casefold())
            except Exception:
                continue
        if not rects:
            continue
        runs: list[tuple[str, float, float]] = []

        try:
            page.extract_text(visitor_text=_make_text_run_collector(runs))
        except Exception:
            runs = []
        for target, rect in rects:
            label = ""
            if rect:
                x0, y0, x1, y1 = rect
                parts = [
                    text
                    for text, x, y in runs
                    if y0 - 2 <= y <= y1 + 2 and x0 - 2 <= x <= x1 + 2
                ]
                label = re.sub(r"\s+", " ", " ".join(parts)).strip()
                if label.casefold() == target.casefold() or len(label) > 60:
                    label = ""
            links.append(f"{label} {target}".strip() if label else target)
    return links


def _append_missing_links(text: str, links: Iterable[str]) -> str:
    additions: list[str] = []
    haystack = text.casefold()
    for link in links:
        target_match = re.search(r"https?://\S+|www\.\S+", link, re.I)
        target = (target_match.group(0) if target_match else link).rstrip(").,;")
        if target.casefold() in haystack:
            continue
        additions.append(link)
        haystack += " " + target.casefold()
    if not additions:
        return text
    return _normalize_extracted_text(f"{text}\n\nEmbedded links\n" + "\n".join(additions))


def _extract_pdf_text(content: bytes) -> str:
    """Extract PDF text via fast backends: PyMuPDF → pdfplumber → pypdf.

    Implementation lives in extractors/pdf.py (no heavy ML document converters).
    """
    from app.features.document_parsing.extractors.pdf import parse_pdf_to_blocks

    blocks = parse_pdf_to_blocks(content)
    text = _normalize_extracted_text(
        "\n".join(block.text for block in blocks if getattr(block, "text", None))
    )
    if not text:
        raise ApiError(422, "document_has_no_text", "No usable text was found in this PDF.")
    if len(text) < MIN_PDF_TEXT_CHARS:
        raise ApiError(
            422,
            "document_has_no_text",
            f"Extracted PDF text is too short to use (need at least {MIN_PDF_TEXT_CHARS} characters).",
        )
    return _append_missing_links(text, _pdf_embedded_links(content))


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        logger.exception("docx_extract_failed")
        raise ApiError(
            400,
            "document_parse_failed",
            "Could not parse this DOCX document.",
        ) from exc
    text = _normalize_extracted_text("\n".join(_docx_paragraph_text(document)))
    if not text:
        raise ApiError(422, "document_has_no_text", "No usable text was found in the DOCX.")
    if len(text) < MIN_DOCX_TEXT_CHARS:
        raise ApiError(
            422,
            "document_has_no_text",
            f"Extracted DOCX text is too short to use (need at least {MIN_DOCX_TEXT_CHARS} characters).",
        )
    return _append_missing_links(text, _docx_embedded_links(document))


def extract_text(content: bytes, mime_type: str) -> str:
    if not content:
        raise ApiError(400, "empty_document", "The selected document is empty.")
    if mime_type == PDF_MIME:
        return _extract_pdf_text(content)
    if mime_type == DOCX_MIME:
        return _extract_docx_text(content)
    raise ApiError(415, "unsupported_document_type", "Only PDF and DOCX documents are supported.")
